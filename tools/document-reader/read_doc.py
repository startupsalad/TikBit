#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word / PDF 读取 —— 抽正文文字（含表格）。

分工：
    .pdf            PyMuPDF 直接抽文字，最快最准
    .docx / .docm   python-docx 直接抽，不需要装 Office
    .doc / .rtf     老二进制格式，python-docx 读不了 —— 走引擎转 PDF 再抽
    .wps            WPS 文字自家格式，同样走引擎

引擎自动降级（详见 engine.py）：Microsoft Office → WPS Office → LibreOffice。

用法：
    python read_doc.py <文件>                  # 打印正文（默认，不产生任何文件）
    python read_doc.py <文件> --maxchars 0     # 不截断，打全文
    python read_doc.py <文件> --maxchars 30000 # 自定义截断字数（默认 12000）
    python read_doc.py <文件> --pdf            # 转 PDF 视觉读（要看版面/图片时）
    python read_doc.py <文件> --out            # 需要存档时才写 <源文件>.txt
    python read_doc.py <文件> --out a.txt      # 写到指定路径（建议放系统临时目录）

注意：默认只打印不落文件。以前每读一次就在源文件旁留个 .txt 副本，
      本库在微盘实时同步，那些副本会到处传染，所以改成必须显式 --out。

      默认截断 12000 字（对齐 read_web.py 的 --limit），**截断会明确告知剩多少字**，
      不是静默丢内容。长标书/合同要通读就加 --maxchars 0。存档（--out）永远写全文。
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))  # 保证能 import engine
import engine

DIRECT_WORD = ('.docx', '.docm')          # python-docx 能直接读的
VIA_ENGINE = ('.doc', '.rtf', '.wps')     # 必须靠 Office/WPS/LO 转换的老格式


def read_pdf(path):
    """PDF 抽文字。PyMuPDF 优先，没装则退 pypdf / PyPDF2。"""
    try:
        import fitz
        with fitz.open(path) as doc:
            return ''.join(f'\n=== 第 {i + 1} 页 ===\n' + p.get_text()
                           for i, p in enumerate(doc))
    except ImportError:
        pass
    try:
        try:
            from pypdf import PdfReader
        except ImportError:
            from PyPDF2 import PdfReader
        reader = PdfReader(path)
        return ''.join(f'\n=== 第 {i + 1} 页 ===\n' + (p.extract_text() or '')
                       for i, p in enumerate(reader.pages))
    except Exception as exc:
        raise RuntimeError(f'PDF 读取失败: {type(exc).__name__}: {exc}')


def read_docx(path):
    """Word 抽正文 + 表格。段落按顺序，表格转成 | 分隔。"""
    from docx import Document
    doc = Document(path)
    out = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        out.append('\n[表格]')
        for row in t.rows:
            out.append(' | '.join(c.text.strip() for c in row.cells))
    return '\n'.join(out)


def read_legacy(path):
    """.doc/.rtf/.wps —— 先转 PDF 再抽文字。这是老格式唯一可靠的路子。"""
    pdf, used = engine.office_to_pdf(path)
    print(f'[{os.path.splitext(path)[1]} 老格式，经 {engine.ENGINE_NAME[used]} '
          f'转 PDF 后抽取]', file=sys.stderr)
    return read_pdf(pdf)


def main():
    engine.fix_console()
    args = list(sys.argv[1:])
    want_pdf = engine.take_flag(args, '--pdf')
    maxchars = engine.take_opt(args, '--maxchars', int, 12000)
    write_out = '--out' in args
    out_path = engine.take_opt(args, '--out')

    if not args:
        print('用法: python read_doc.py <文件> [--pdf] [--maxchars 12000] '
              '[--out [输出路径]]')
        sys.exit(1)

    src = args[0]
    if not os.path.exists(src):
        print(f'文件不存在: {src}')
        print('排查：微盘"按需下载"的占位符读不了，先在资源管理器里双击下载到本地。')
        sys.exit(1)

    ext = os.path.splitext(src)[1].lower()

    if want_pdf:      # 要看版面/图片，转 PDF 视觉读
        if ext == '.pdf':
            print(f'本来就是 PDF，直接 Read 即可：\n{os.path.abspath(src)}')
            return
        try:
            pdf, used = engine.office_to_pdf(src, out=out_path)
            print(f'PDF 已生成（引擎 {engine.ENGINE_NAME[used]}）：\n{pdf}')
            print('→ AI 直接 Read 这个 PDF 即可视觉读取。')
        except Exception as exc:
            print(f'转换失败: {exc}')
            sys.exit(1)
        return

    try:
        if ext == '.pdf':
            content = read_pdf(src)
        elif ext in DIRECT_WORD:
            content = read_docx(src)
        elif ext in VIA_ENGINE:
            content = read_legacy(src)
        else:
            print(f'不支持的格式: {ext}')
            print('本工具处理 .pdf / .docx / .docm / .doc / .rtf / .wps；'
                  '表格走 read_excel.py，PPT 走 read_ppt.py，图片走 read_img.py。')
            sys.exit(1)
    except Exception as exc:
        print(f'读取失败: {type(exc).__name__}: {exc}')
        sys.exit(1)

    if write_out:      # 存档写全文，不截断——归档里不留残缺副本
        dst = out_path or (src + '.txt')
        with open(dst, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f'内容已保存到: {dst}（全文 {len(content)} 字）')
    else:
        print(engine.clip(content, maxchars))


if __name__ == '__main__':
    main()
